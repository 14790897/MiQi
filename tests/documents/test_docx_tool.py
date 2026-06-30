"""Tests for miqi.documents.docx_tool."""

import tempfile
from pathlib import Path


def test_imports():
    from miqi.documents.docx_tool import DocxReadTool, DocxWriteTool  # noqa: F401


def test_docx_read_file_not_found():
    import asyncio
    from miqi.documents.docx_tool import DocxReadTool

    tool = DocxReadTool()
    result = asyncio.run(tool.execute(file_path="/nonexistent/doc.docx"))
    assert "Error" in result
    assert "not found" in result


def test_docx_read_valid_file():
    import asyncio
    from miqi.documents.docx_tool import DocxReadTool
    from docx import Document

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph")
        doc.save(str(path))

        tool = DocxReadTool()
        result = asyncio.run(tool.execute(file_path=str(path)))
        assert "Hello world" in result
        assert "Second paragraph" in result


def test_docx_write_creates_file():
    import asyncio
    from miqi.documents.docx_tool import DocxWriteTool

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "output.docx"
        tool = DocxWriteTool()
        result = asyncio.run(tool.execute(
            file_path=str(path),
            content="# Title\n\nParagraph one.\n\nParagraph two.",
        ))
        assert path.exists()
        assert "created" in result.lower() or "ok" in result.lower()

        # Verify content
        from docx import Document
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Title" in texts
        assert "Paragraph one." in texts
