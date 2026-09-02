"""Structured preview parsing (issue #877): xlsx/csv spreadsheets + docx blocks.

These tests cover the `structured=True` payloads the frontend rich preview
(XLSX/CSV table, DOCX document render) consumes.
"""

from __future__ import annotations

import csv
import tempfile
from pathlib import Path

from miqi.documents.document_parser import parse_document


def test_xlsx_structured_sheets_rows_and_merges():
    import openpyxl

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "params.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "参数"
        ws["B1"] = "值"
        ws["A2"] = "温度"
        ws["B2"] = 300
        # merged block approximating a section header spanning A3:B3
        ws.merge_cells("A3:B3")
        ws["A3"] = "合并区"
        wb.create_sheet("Sheet2")
        wb.save(str(path))

        result = parse_document(path, structured=True)

        assert result["structured"]["kind"] == "spreadsheet"
        sheets = result["structured"]["sheets"]
        assert [s["name"] for s in sheets] == ["Sheet1", "Sheet2"]
        sheet1 = sheets[0]
        assert sheet1["rows"][0] == ["参数", "值"]
        assert sheet1["rows"][1] == ["温度", "300"]
        assert sheet1["rows"][2] == ["合并区", ""]
        assert sheet1["merges"] == [
            {"start_row": 2, "start_col": 0, "end_row": 2, "end_col": 1}
        ]


def test_xlsx_structured_absent_without_flag():
    import openpyxl

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "plain.xlsx"
        wb = openpyxl.Workbook()
        wb.active["A1"] = "hello"
        wb.save(str(path))

        result = parse_document(path)
        assert "structured" not in result
        assert "hello" in result["text"]


def test_csv_structured_single_sheet():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "data.csv"
        with path.open("w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["name", "count"])
            writer.writerow(["a", "1"])
            writer.writerow(["b", "2"])

        result = parse_document(path, structured=True)

        assert result["structured"]["kind"] == "spreadsheet"
        sheets = result["structured"]["sheets"]
        assert len(sheets) == 1
        assert sheets[0]["name"] == "data"
        assert sheets[0]["rows"] == [["name", "count"], ["a", "1"], ["b", "2"]]
        assert "merges" not in sheets[0]


def test_docx_structured_ordered_blocks_with_table_and_image():
    import struct
    import zlib

    from docx import Document

    def make_png() -> bytes:
        """Programmatic 1x1 red RGBA PNG (valid IHDR/IDAT/IEND)."""
        def chunk(tag: bytes, data: bytes) -> bytes:
            body = tag + data
            return (
                struct.pack(">I", len(data))
                + body
                + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)
            )

        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 6, 0, 0, 0)
        raw = b"\x00\xff\x00\x00"  # filter byte 0 + red RGBA pixel
        return (
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", zlib.compress(raw))
            + chunk(b"IEND", b"")
        )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "report.docx"
        doc = Document()
        doc.add_heading("标题", level=1)
        doc.add_paragraph("正文段落")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "H1"
        table.cell(0, 1).text = "H2"
        table.cell(1, 0).text = "v1"
        table.cell(1, 1).text = "v2"
        with Path(tmp, "p.png").open("wb") as f:
            f.write(make_png())
        doc.add_picture(str(Path(tmp, "p.png")))
        doc.save(str(path))

        result = parse_document(path, structured=True)

        blocks = result["structured"]["blocks"]
        assert result["structured"]["kind"] == "document"
        kinds = [b["type"] for b in blocks]
        assert kinds[0] == "heading"
        assert blocks[0]["level"] == 1
        assert blocks[0]["text"] == "标题"
        assert "paragraph" in kinds
        table_block = next(b for b in blocks if b["type"] == "table")
        assert table_block["rows"] == [["H1", "H2"], ["v1", "v2"]]
        image_block = next(b for b in blocks if b["type"] == "image")
        assert image_block["data_url"].startswith("data:image/png;base64,")


def test_docx_structured_absent_when_python_docx_missing_doc():
    """.doc (binary Word) has no structured support — falls back to text."""
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "legacy.doc"
        path.write_bytes(b"\xd0\xcf\x11\xe0 not a real doc")

        result = parse_document(path, structured=True)
        assert "structured" not in result
        assert "解析失败" in result["text"]


def test_docx_structured_caps_block_text_and_total_budget():
    """CWE-400 guard (#889): block text capped + aggregate text budget."""
    from docx import Document

    from miqi.documents.document_parser import (
        _MAX_STRUCTURED_BLOCK_TEXT_CHARS,
        _MAX_STRUCTURED_TEXT_TOTAL,
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "big.docx"
        doc = Document()
        # 60 paragraphs × 4000 chars — far beyond the aggregate budget
        for i in range(60):
            doc.add_paragraph(f"段落{i} " + "x" * 4000)
        doc.save(str(path))

        result = parse_document(path, structured=True)

        blocks = result["structured"]["blocks"]
        text_blocks = [b for b in blocks if b["type"] in ("heading", "paragraph")]
        assert text_blocks, "some text blocks should be emitted"
        for b in text_blocks:
            assert len(b["text"]) <= _MAX_STRUCTURED_BLOCK_TEXT_CHARS
        total = sum(len(b["text"]) for b in text_blocks)
        # The budget check runs BEFORE appending, so one block may overshoot;
        # allow a single block's slack on top of the budget.
        assert total <= _MAX_STRUCTURED_TEXT_TOTAL + _MAX_STRUCTURED_BLOCK_TEXT_CHARS
        # The 60 paragraphs must NOT all be present (budget cut the tail off)
        assert len(text_blocks) < 60


def test_parse_handler_rejects_oversized_inline_attachment():
    """CWE-400 guard (#889): oversized data_base64 rejected before decoding."""
    import asyncio

    import pytest

    from miqi.documents.documents_parse_handler import (
        _MAX_INLINE_ATTACHMENT_B64_CHARS,
        documents_parse_handler,
    )
    from miqi.runtime.app_server import AppServerError

    huge = "A" * (_MAX_INLINE_ATTACHMENT_B64_CHARS + 1)
    with pytest.raises(AppServerError, match="too large"):
        asyncio.run(
            documents_parse_handler(
                "req-test", {"path": "x.xlsx", "data_base64": huge},
                "client", None, None,
            )
        )
