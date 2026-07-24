"""Word document (.docx) read/write tools."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any

from miqi.agent.tools.base import Tool
from miqi.agent.tools.filesystem import _persist_tracked_file


_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")

_CHINESE_STYLE_PRESETS: dict[str, dict[str, dict[str, Any]]] = {
    "chinese_document": {
        "title_style": {
            "font_name": "黑体",
            "east_asia_font": "黑体",
            "font_size_pt": 16,
            "bold": True,
            "alignment": "center",
        },
        "body_style": {
            "font_name": "宋体",
            "east_asia_font": "宋体",
            "font_size_pt": 12,
            "line_spacing": 1.5,
        },
    },
    "chinese_essay": {
        "title_style": {
            "font_name": "黑体",
            "east_asia_font": "黑体",
            "font_size_pt": 16,
            "bold": True,
            "alignment": "center",
        },
        "body_style": {
            "font_name": "宋体",
            "east_asia_font": "宋体",
            "font_size_pt": 12,
            "line_spacing": 1.5,
        },
    },
}

_CHINESE_SIZE_TO_PT = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
}


def _raw_output_path(kwargs: dict[str, Any]) -> str:
    return str(
        kwargs.get("filename")
        or kwargs.get("file_path")
        or kwargs.get("path")
        or ""
    )


def _ensure_suffix(path: Path, suffix: str) -> Path:
    if not path.name or path.name in {".", ".."}:
        raise ValueError("output filename is required")
    if path.suffix.lower() == suffix:
        return path
    return path.with_suffix(suffix)


def _enforce_boundary(path: Path, allowed_dir: Path | None, workspace: Path | None) -> None:
    effective_dir = allowed_dir or workspace
    if effective_dir is None:
        return
    try:
        path.resolve().relative_to(effective_dir.resolve())
    except ValueError:
        raise PermissionError(
            f"Path '{path}' resolves outside allowed directory '{effective_dir}'"
        )


def _add_docx_content(doc: Any, content: Any) -> int:
    """Add supported structured content to a python-docx document."""
    blocks = content if isinstance(content, list) else [{"type": "paragraph", "text": str(content)}]
    count = 0
    for block in blocks:
        if not isinstance(block, dict):
            doc.add_paragraph(str(block))
            count += 1
            continue
        block_type = str(block.get("type", "paragraph")).lower()
        if block_type == "heading":
            doc.add_heading(str(block.get("text", "")), level=int(block.get("level", 1)))
            count += 1
        elif block_type == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = str(block.get("style", "Table Grid"))
            for row_idx, row in enumerate(rows):
                for col_idx, value in enumerate(row):
                    table.cell(row_idx, col_idx).text = "" if value is None else str(value)
            count += 1
        elif block_type == "image":
            image_path = block.get("path")
            if image_path:
                doc.add_picture(str(image_path))
                count += 1
        else:
            doc.add_paragraph(str(block.get("text", "")))
            count += 1
    return count


def _add_markdown_like_text(doc: Any, text: Any) -> int:
    count = 0
    for line in str(text).split("\n"):
        line = line.strip()
        if not line:
            continue
        heading = _MARKDOWN_HEADING_RE.match(line)
        if heading:
            doc.add_heading(heading.group(2), level=len(heading.group(1)))
        else:
            doc.add_paragraph(line)
        count += 1
    return count


def _set_east_asia_font(run: Any, font_name: str) -> None:
    from docx.oxml.ns import qn

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def _alignment_from_value(value: Any) -> Any:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if value is None:
        return None
    normalized = str(value).strip().lower()
    return {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "centered": WD_ALIGN_PARAGRAPH.CENTER,
        "centre": WD_ALIGN_PARAGRAPH.CENTER,
        "居中": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(normalized)


def _coerce_bool(value: Any) -> bool | None:
    if value is None or isinstance(value, bool):
        return value
    normalized = str(value).strip().lower()
    if normalized in {"true", "1", "yes", "y", "bold", "加粗"}:
        return True
    if normalized in {"false", "0", "no", "n", "normal", "不加粗"}:
        return False
    return None


def _size_to_pt(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, int | float):
        return float(value)
    text = str(value).strip()
    if text in _CHINESE_SIZE_TO_PT:
        return float(_CHINESE_SIZE_TO_PT[text])
    try:
        return float(text.replace("pt", "").strip())
    except ValueError:
        return None


def _merge_style(*styles: dict[str, Any] | None) -> dict[str, Any]:
    merged: dict[str, Any] = {}
    for style in styles:
        if isinstance(style, dict):
            merged.update({k: v for k, v in style.items() if v is not None})
    return merged


def _instruction_clause(instructions: str, label: str) -> str:
    start = instructions.find(label)
    if start < 0:
        return ""
    other_labels = ["标题", "正文"]
    end = len(instructions)
    for other in other_labels:
        if other == label:
            continue
        other_pos = instructions.find(other, start + len(label))
        if other_pos >= 0:
            end = min(end, other_pos)
    return instructions[start:end]


def _size_in_clause(clause: str) -> float | None:
    # Match longer names first so 小四 is not treated as 四号, etc.
    for chinese_size in sorted(_CHINESE_SIZE_TO_PT, key=len, reverse=True):
        if chinese_size in clause:
            return float(_CHINESE_SIZE_TO_PT[chinese_size])
    return None


def _style_from_kwargs(kwargs: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    preset_name = str(kwargs.get("style_preset") or "").strip()
    preset = _CHINESE_STYLE_PRESETS.get(preset_name, {})
    title_style = _merge_style(preset.get("title_style"), kwargs.get("title_style"))
    body_style = _merge_style(preset.get("body_style"), kwargs.get("body_style"))

    flat_title = {
        "font_name": kwargs.get("title_font_name") or kwargs.get("title_font"),
        "east_asia_font": kwargs.get("title_east_asia_font"),
        "font_size_pt": kwargs.get("title_font_size_pt") or kwargs.get("title_size_pt"),
        "bold": kwargs.get("title_bold"),
        "alignment": kwargs.get("title_alignment"),
    }
    flat_body = {
        "font_name": kwargs.get("body_font_name") or kwargs.get("body_font"),
        "east_asia_font": kwargs.get("body_east_asia_font"),
        "font_size_pt": kwargs.get("body_font_size_pt") or kwargs.get("body_size_pt"),
        "line_spacing": kwargs.get("line_spacing"),
        "alignment": kwargs.get("body_alignment"),
    }
    title_style = _merge_style(title_style, flat_title)
    body_style = _merge_style(body_style, flat_body)

    instructions = str(kwargs.get("format_instructions") or "")
    if instructions:
        title_clause = _instruction_clause(instructions, "标题")
        body_clause = _instruction_clause(instructions, "正文")
        if "黑体" in instructions:
            title_style["font_name"] = "黑体"
            title_style["east_asia_font"] = "黑体"
        if "宋体" in instructions:
            body_style["font_name"] = "宋体"
            body_style["east_asia_font"] = "宋体"
        title_size = _size_in_clause(title_clause)
        body_size = _size_in_clause(body_clause)
        if title_size is not None:
            title_style["font_size_pt"] = title_size
        if body_size is not None:
            body_style["font_size_pt"] = body_size
        if "1.5" in instructions or "1.5倍" in instructions or "1.5行距" in instructions:
            body_style["line_spacing"] = 1.5
        if "居中" in instructions:
            title_style["alignment"] = "center"
        if "加粗" in instructions:
            title_style["bold"] = True

    return title_style, body_style


def _apply_run_style(run: Any, style: dict[str, Any]) -> None:
    from docx.shared import Pt

    font_name = style.get("font_name")
    if font_name:
        run.font.name = str(font_name)
    east_asia_font = style.get("east_asia_font") or font_name
    if east_asia_font:
        _set_east_asia_font(run, str(east_asia_font))
    size_pt = _size_to_pt(style.get("font_size_pt") or style.get("size"))
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    bold = _coerce_bool(style.get("bold"))
    if bold is not None:
        run.bold = bold


def _apply_paragraph_style(paragraph: Any, style: dict[str, Any]) -> None:
    if not style:
        return
    alignment = _alignment_from_value(style.get("alignment"))
    if alignment is not None:
        paragraph.alignment = alignment
    if style.get("line_spacing") is not None:
        try:
            paragraph.paragraph_format.line_spacing = float(style["line_spacing"])
        except (TypeError, ValueError):
            pass
    for run in paragraph.runs:
        _apply_run_style(run, style)


def _is_heading_paragraph(paragraph: Any) -> bool:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    return style_name.startswith("Heading") or style_name == "Title"


def _iter_all_paragraphs(doc: Any) -> Any:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _apply_document_styles(
    doc: Any,
    title_style: dict[str, Any] | None = None,
    body_style: dict[str, Any] | None = None,
) -> None:
    title_style = title_style or {}
    body_style = body_style or {}
    seen_title = False
    for paragraph in _iter_all_paragraphs(doc):
        if not paragraph.text.strip():
            continue
        if _is_heading_paragraph(paragraph):
            if title_style and not seen_title:
                _apply_paragraph_style(paragraph, title_style)
                seen_title = True
            continue
        if body_style:
            _apply_paragraph_style(paragraph, body_style)


class DocxReadTool(Tool):
    """Read the text content of a Word (.docx) document."""

    name = "docx_read"
    description = "Read and extract text content from a Word (.docx) document."

    def __init__(
        self,
        workspace: Path | None = None,
        allowed_dir: Path | None = None,
    ):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the .docx file to read. Alias for filename.",
                },
                "filename": {
                    "type": "string",
                    "description": "Filename or relative path to the .docx file.",
                },
                "path": {
                    "type": "string",
                    "description": "Path to the .docx file to read. Alias for filename.",
                },
            },
            "anyOf": [
                {"required": ["filename"]},
                {"required": ["file_path"]},
                {"required": ["path"]},
            ],
        }

    async def execute(self, **kwargs: Any) -> str:
        raw_path = _raw_output_path(kwargs)
        if not raw_path.strip():
            return "Error: filename is required"
        try:
            file_path = _resolve_output_path(
                raw_path, self._workspace, self._allowed_dir,
            )
            file_path = _ensure_suffix(file_path, ".docx")
            _enforce_boundary(file_path, self._allowed_dir, self._workspace)
        except PermissionError as e:
            return f"Error: Permission denied: {e}"
        except ValueError as e:
            return f"Error: {e}"
        if not file_path.exists():
            return f"Error: file not found: {file_path}"
        try:
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            return "\n\n".join(paragraphs)
        except Exception as e:
            return f"Error reading {file_path.name}: {e}"


def _resolve_output_path(
    file_path: str,
    workspace: Path | None,
    allowed_dir: Path | None,
) -> Path:
    """Resolve an output path and enforce workspace/directory bounds.

    Office document write tools always write inside the workspace:
    - Relative paths are resolved against *workspace*.
    - If *allowed_dir* is ``None`` but *workspace* is set, *workspace*
      is used as the effective boundary (defense-in-depth default).
    - Absolute paths outside the effective boundary are rejected.

    Raises:
        PermissionError: if the resolved path is outside the effective boundary.
    """
    p = Path(file_path).expanduser()
    if not p.is_absolute() and workspace is not None:
        p = workspace / p
    resolved = p.resolve()

    # Defense-in-depth: when no explicit allowed_dir is given, office
    # write tools default to workspace as the boundary.  This is
    # independent of the `restrict_to_workspace` config (which only
    # controls WriteFileTool / EditFileTool).
    effective_dir = allowed_dir
    if effective_dir is None and workspace is not None:
        effective_dir = workspace.resolve()

    if effective_dir is not None:
        try:
            resolved.relative_to(effective_dir.resolve())
        except ValueError:
            raise PermissionError(
                f"Path '{file_path}' resolves outside allowed directory "
                f"'{effective_dir}'"
            )
    return resolved


class CreateDocxTool(Tool):
    """Create or overwrite a Word (.docx) document."""

    name = "create_docx"
    description = (
        "Create a Word (.docx) document in the workspace files directory. "
        "Supports title, paragraphs, headings, tables, images, and common "
        "Word formatting such as Chinese fonts, font sizes, alignment, bold, "
        "and line spacing."
    )

    def __init__(
        self,
        workspace: Path | None = None,
        allowed_dir: Path | None = None,
    ):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path for the output .docx file. Alias for filename.",
                },
                "filename": {
                    "type": "string",
                    "description": "Filename or relative path for the output .docx file",
                },
                "path": {
                    "type": "string",
                    "description": "Path for the output .docx file. Alias for filename.",
                },
                "title": {
                    "type": "string",
                    "description": "Optional document title",
                },
                "content": {
                    "description": (
                        "Document content. Use a string for markdown-like text, "
                        "or an array of blocks: {type: paragraph|heading|table|image, ...}."
                    ),
                },
                "paragraphs": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional paragraph list",
                },
                "tables": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "rows": {"type": "array", "items": {"type": "array"}},
                            "style": {"type": "string"},
                        },
                        "required": ["rows"],
                    },
                    "description": "Optional tables to append",
                },
                "style_preset": {
                    "type": "string",
                    "enum": ["chinese_document", "chinese_essay"],
                    "description": (
                        "Optional formatting preset. Use chinese_document or "
                        "chinese_essay for Chinese Word requirements such as "
                        "标题黑体三号加粗居中、正文宋体小四、1.5 行距."
                    ),
                },
                "title_style": {
                    "type": "object",
                    "description": (
                        "Formatting for the main title. Supports font_name, "
                        "east_asia_font, font_size_pt, bold, alignment. "
                        "Use for requests like 标题黑体加粗三号字居中."
                    ),
                },
                "body_style": {
                    "type": "object",
                    "description": (
                        "Formatting for body paragraphs and table text. "
                        "Supports font_name, east_asia_font, font_size_pt, "
                        "line_spacing, alignment. Use for requests like "
                        "正文宋体小四、段落 1.5 行距."
                    ),
                },
                "format_instructions": {
                    "type": "string",
                    "description": (
                        "Natural language formatting instructions to apply, "
                        "for example: 正文宋体小四，段落1.5行距，标题黑体加粗三号字居中."
                    ),
                },
            },
            "anyOf": [
                {"required": ["filename"]},
                {"required": ["file_path"]},
                {"required": ["path"]},
            ],
        }

    async def execute(self, **kwargs: Any) -> str:
        from docx import Document

        _sess_key = kwargs.pop("_session_key", None)
        raw_path = _raw_output_path(kwargs)
        content = kwargs.get("content", "")
        if not raw_path.strip():
            return "Error: filename is required"

        try:
            file_path = _resolve_output_path(
                raw_path, self._workspace, self._allowed_dir,
            )
            file_path = _ensure_suffix(file_path, ".docx")
            _enforce_boundary(file_path, self._allowed_dir, self._workspace)
        except PermissionError as e:
            return f"Error: Permission denied: {e}"
        except ValueError as e:
            return f"Error: {e}"
        if not (
            kwargs.get("title")
            or content
            or kwargs.get("paragraphs")
            or kwargs.get("tables")
        ):
            return "Error: provide title, content, paragraphs, or tables"

        try:
            doc = Document()
            if kwargs.get("title"):
                doc.add_heading(str(kwargs["title"]), level=0)
            if isinstance(content, list):
                _add_docx_content(doc, content)
            elif content:
                _add_markdown_like_text(doc, content)
            for paragraph in kwargs.get("paragraphs", []) or []:
                doc.add_paragraph(str(paragraph))
            for table_data in kwargs.get("tables", []) or []:
                _add_docx_content(doc, [{"type": "table", **table_data}])

            title_style, body_style = _style_from_kwargs(kwargs)
            _apply_document_styles(doc, title_style, body_style)

            file_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(file_path))
            _persist_tracked_file(self._workspace, file_path, op="write", session_key=_sess_key)
            return f"Created: {file_path}"
        except Exception as e:
            return f"Error writing {raw_path}: {e}"


class DocxWriteTool(CreateDocxTool):
    """Backward-compatible alias for create_docx."""

    name = "docx_write"
    description = (
        "Create a new Word (.docx) document with the given content. "
        "Prefer create_docx for new calls."
    )


class EditDocxTool(Tool):
    """Edit an existing Word (.docx) document."""

    name = "edit_docx"
    description = (
        "Edit an existing Word (.docx) document in the workspace files directory. "
        "Supports text replacement, appending paragraphs, and applying common "
        "Word formatting such as Chinese fonts, font sizes, alignment, bold, "
        "and line spacing."
    )

    def __init__(
        self,
        workspace: Path | None = None,
        allowed_dir: Path | None = None,
    ):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the .docx file. Alias for filename.",
                },
                "filename": {
                    "type": "string",
                    "description": "Filename or relative path to the .docx file.",
                },
                "path": {
                    "type": "string",
                    "description": "Path to the .docx file. Alias for filename.",
                },
                "old_text": {
                    "type": "string",
                    "description": "Text to replace.",
                },
                "new_text": {
                    "type": "string",
                    "description": "Replacement text.",
                },
                "append_paragraphs": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Paragraphs to append to the end of the document.",
                },
                "content": {
                    "type": "string",
                    "description": "Text content to append as paragraphs.",
                },
                "style_preset": {
                    "type": "string",
                    "enum": ["chinese_document", "chinese_essay"],
                    "description": (
                        "Optional formatting preset. Use chinese_document or "
                        "chinese_essay for Chinese Word requirements such as "
                        "标题黑体三号加粗居中、正文宋体小四、1.5 行距."
                    ),
                },
                "title_style": {
                    "type": "object",
                    "description": (
                        "Formatting for the main title. Supports font_name, "
                        "east_asia_font, font_size_pt, bold, alignment. "
                        "Use for requests like 标题黑体加粗三号字居中."
                    ),
                },
                "body_style": {
                    "type": "object",
                    "description": (
                        "Formatting for body paragraphs and table text. "
                        "Supports font_name, east_asia_font, font_size_pt, "
                        "line_spacing, alignment. Use for requests like "
                        "正文宋体小四、段落 1.5 行距."
                    ),
                },
                "format_instructions": {
                    "type": "string",
                    "description": (
                        "Natural language formatting instructions to apply, "
                        "for example: 正文宋体小四，段落1.5行距，标题黑体加粗三号字居中."
                    ),
                },
            },
            "anyOf": [
                {"required": ["filename"]},
                {"required": ["file_path"]},
                {"required": ["path"]},
            ],
        }

    async def execute(self, **kwargs: Any) -> str:
        from docx import Document

        _sess_key = kwargs.pop("_session_key", None)
        raw_path = _raw_output_path(kwargs)
        if not raw_path.strip():
            return "Error: filename is required"

        try:
            file_path = _resolve_output_path(
                raw_path, self._workspace, self._allowed_dir,
            )
            file_path = _ensure_suffix(file_path, ".docx")
            _enforce_boundary(file_path, self._allowed_dir, self._workspace)
        except PermissionError as e:
            return f"Error: Permission denied: {e}"
        except ValueError as e:
            return f"Error: {e}"

        if not file_path.exists():
            return f"Error: file not found: {file_path}"

        old_text = kwargs.get("old_text")
        new_text = kwargs.get("new_text")
        append_paragraphs = list(kwargs.get("append_paragraphs") or [])
        if kwargs.get("content"):
            append_paragraphs.extend(
                line.strip()
                for line in str(kwargs["content"]).splitlines()
                if line.strip()
            )
        title_style, body_style = _style_from_kwargs(kwargs)
        has_formatting = bool(title_style or body_style)

        if not (old_text and new_text is not None) and not append_paragraphs and not has_formatting:
            return (
                "Error: provide old_text/new_text, append_paragraphs/content, "
                "or formatting instructions"
            )

        try:
            doc = Document(str(file_path))
            replacements = 0
            if old_text and new_text is not None:
                for paragraph in doc.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, str(new_text))
                        replacements += 1
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if old_text in paragraph.text:
                                    paragraph.text = paragraph.text.replace(old_text, str(new_text))
                                    replacements += 1
                if replacements == 0:
                    return f"Error: old_text not found in {file_path}"

            for paragraph in append_paragraphs:
                doc.add_paragraph(str(paragraph))

            if has_formatting:
                _apply_document_styles(doc, title_style, body_style)

            doc.save(str(file_path))
            return (
                f"Edited: {file_path} "
                f"({replacements} replacement(s), {len(append_paragraphs)} appended"
                f"{', formatted' if has_formatting else ''})"
            )
        except Exception as e:
            return f"Error editing {raw_path}: {e}"
